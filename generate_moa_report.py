"""
generate_moa_report.py
──────────────────────
Reads scored data from BigQuery `moa_innovation_table` and generates a
professional Word (.docx) report using Gemini for narrative generation.

Only the LATEST row per drug (by created_at) is used.

Usage:
    python market_potential/generate_moa_report.py
    python market_potential/generate_moa_report.py --molecule Semaglutide
    python market_potential/generate_moa_report.py --output moa_report.docx
"""

import os
import re
import json
import argparse
from datetime import datetime
from pathlib import Path

try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

import pandas as pd
from google.cloud import bigquery
from google.oauth2 import service_account
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from google import genai as genai_client
from google.genai import types

# ── Config ────────────────────────────────────────────────────────────────────
API_KEY = os.environ.get("GEMINI_API_KEY", "")
CREDENTIALS_PATH = os.environ.get("GOOGLE_APPLICATION_CREDENTIALS", "")
MODEL = "gemini-2.5-flash"

BQ_PROJECT_ID = os.environ.get("PROJECT_ID", "cognito-prod-394707")
BQ_DATASET_ID = os.environ.get("BQ_DATASET_ID", "cognito_prod_datamart")
BQ_LOCATION = "asia-south1"
BQ_TABLE = "moa_innovation_table"

GCS_BUCKET = "cognito-gcs"
GCS_BASE_PATH = "Cognito_new/reports"
GCS_FILE_NAME = "MoA_Innovation_Analysis.docx"

REPORT_TITLE = "MoA Innovation Scoring Analysis"

SCORE_COLOR_MAP = {
    5: RGBColor(0x00, 0x80, 0x00),  # Green  – Exceptional
    4: RGBColor(0x4C, 0xAF, 0x50),  # Light green – Strong
    3: RGBColor(0xCC, 0x99, 0x00),  # Amber – Moderate
    2: RGBColor(0xE6, 0x51, 0x00),  # Orange-red – Weak
    1: RGBColor(0xCC, 0x00, 0x00),  # Red – Poor
}

SCORE_LABEL = {
    5: "Exceptional",
    4: "Strong",
    3: "Moderate",
    2: "Weak",
    1: "Poor",
}

CLASSIFICATION_COLORS = {
    "First-in-Class": RGBColor(0x00, 0x80, 0x00),
    "Best-in-Class": RGBColor(0x4C, 0xAF, 0x50),
    "Me-too": RGBColor(0xCC, 0x99, 0x00),
    "Weak/Outdated": RGBColor(0xE6, 0x51, 0x00),
    "Poor/Invalid": RGBColor(0xCC, 0x00, 0x00),
}


# ── Gemini helper ─────────────────────────────────────────────────────────────

def call_gemini(prompt: str) -> str:
    client = genai_client.Client(api_key=API_KEY)
    config = types.GenerateContentConfig(temperature=0.3)
    response = client.models.generate_content(model=MODEL, contents=prompt, config=config)
    return response.text.strip() if response.text else ""


def _extract_json(text: str):
    text = re.sub(r"^```(?:json)?", "", text.strip()).strip()
    text = re.sub(r"```$", "", text).strip()
    match = re.search(r"\{.*\}", text, re.DOTALL)
    if match:
        try:
            return json.loads(match.group(0))
        except json.JSONDecodeError:
            pass
    return None


# ── Data loading ──────────────────────────────────────────────────────────────

def _get_credentials():
    if CREDENTIALS_PATH and os.path.exists(CREDENTIALS_PATH):
        return service_account.Credentials.from_service_account_file(CREDENTIALS_PATH)
    return None


def _bq_client():
    credentials = _get_credentials()
    return bigquery.Client(project=BQ_PROJECT_ID, credentials=credentials, location=BQ_LOCATION)


def load_from_bigquery(molecule: str = None) -> pd.DataFrame:
    """Load LATEST row per drug from moa_innovation_table."""
    client = _bq_client()

    molecule_filter = ""
    if molecule:
        molecule_filter = f"AND drug_name = '{molecule}'"

    query = f"""
    WITH ranked AS (
        SELECT *,
            ROW_NUMBER() OVER (PARTITION BY drug_name ORDER BY created_at DESC) AS _rn
        FROM `{BQ_PROJECT_ID}.{BQ_DATASET_ID}.{BQ_TABLE}`
        WHERE drug_name IS NOT NULL {molecule_filter}
    )
    SELECT * EXCEPT(_rn) FROM ranked WHERE _rn = 1
    ORDER BY drug_name
    """
    print(f"Loading latest MoA data from {BQ_TABLE}...")
    df = client.query(query).to_dataframe()
    print(f"  Loaded {len(df)} drug(s).")
    return df


# ── Statistics ────────────────────────────────────────────────────────────────

def compute_statistics(df: pd.DataFrame) -> dict:
    stats = {
        "total_drugs": len(df),
        "drugs": [],
        "avg_score": None,
        "score_distribution": {},
        "classification_distribution": {},
        "guardrail_pass": 0,
        "guardrail_fail": 0,
        "per_drug_stats": {},
    }
    if df.empty:
        return stats

    stats["drugs"] = df["drug_name"].dropna().unique().tolist()

    if "score" in df.columns:
        numeric = pd.to_numeric(df["score"], errors="coerce").dropna()
        if len(numeric):
            stats["avg_score"] = float(round(numeric.mean(), 2))
            for s in range(1, 6):
                stats["score_distribution"][s] = int((numeric == s).sum())

    if "moa_classification" in df.columns:
        for cls_val, cnt in df["moa_classification"].value_counts().items():
            stats["classification_distribution"][str(cls_val)] = int(cnt)

    if "guardrail" in df.columns:
        stats["guardrail_pass"] = int((df["guardrail"] == "PASS").sum())
        stats["guardrail_fail"] = int((df["guardrail"] == "FAIL").sum())

    for _, row in df.iterrows():
        drug = str(row.get("drug_name", ""))
        stats["per_drug_stats"][drug] = {
            "score": row.get("score"),
            "classification": str(row.get("moa_classification", "N/A")),
            "guardrail": str(row.get("guardrail", "N/A")),
            "confidence_tier": str(row.get("confidence_tier", "N/A")),
            "confidence_score": row.get("confidence_score"),
            "indication": str(row.get("indication", "N/A")),
        }

    return stats


# ── LLM narrative ─────────────────────────────────────────────────────────────

def generate_executive_summary(stats: dict, df: pd.DataFrame) -> dict:
    drug_rows = []
    for _, r in df.iterrows():
        drug_rows.append({
            "drug": str(r.get("drug_name", "")),
            "indication": str(r.get("indication", "")),
            "mechanism": str(r.get("mechanism_statement", ""))[:300],
            "classification": str(r.get("moa_classification", "")),
            "score": r.get("score"),
            "guardrail": str(r.get("guardrail", "")),
            "confidence": r.get("confidence_score"),
            "confidence_tier": str(r.get("confidence_tier", "")),
        })

    prompt = f"""You are a senior pharmaceutical analyst writing a MoA Innovation report.
Based on the data below, produce a thorough analytical report. Be specific — reference
drug names, classifications, scores, and indications.

PORTFOLIO STATISTICS:
- Total drugs assessed: {stats['total_drugs']}
- Drugs: {', '.join(stats['drugs'])}
- Average MoA score: {stats['avg_score']} (1=Poor, 5=Exceptional)
- Score distribution: {json.dumps(stats['score_distribution'])}
- Classification distribution: {json.dumps(stats['classification_distribution'])}
- Guardrail: {stats['guardrail_pass']} PASS, {stats['guardrail_fail']} FAIL

DRUG-LEVEL DATA (JSON):
{json.dumps(drug_rows, indent=1)}

Respond ONLY with a valid JSON object (no markdown fences):
{{
  "executive_summary": "<5-8 sentences overview of the MoA innovation landscape across the assessed drugs>",
  "key_findings": ["<finding 1 with drug names and scores>", "<finding 2>", "<finding 3>", "<finding 4>"],
  "classification_analysis": "<4-6 sentences analysing the classification distribution — which drugs are FIC/BIC/Me-too and why>",
  "risk_highlights": "<3-5 sentences on drugs with low scores or FAIL guardrails>",
  "strength_highlights": "<3-5 sentences on drugs with high scores and strong innovation>",
  "per_drug_narratives": {{
    "<drug_name>": "<3-5 sentence analysis of this drug's MoA innovation, classification rationale, and strategic implications>"
  }}
}}
"""
    text = call_gemini(prompt)
    result = _extract_json(text)
    if result:
        return result
    return {
        "executive_summary": "Analysis complete. See tables below for details.",
        "key_findings": ["See detailed tables."],
        "classification_analysis": "See classification table.",
        "risk_highlights": "Refer to score tables.",
        "strength_highlights": "Refer to score tables.",
        "per_drug_narratives": {},
    }


# ── Document builder helpers ──────────────────────────────────────────────────

def set_cell_shading(cell, hex_color: str):
    shading = cell._element.get_or_add_tcPr()
    shading_el = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): hex_color, qn("w:val"): "clear",
    })
    shading.append(shading_el)


def _styled_heading(doc, text, level=2, color=RGBColor(0x1F, 0x38, 0x64), size=Pt(13)):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
        run.font.size = size
    return h


# ── Report builder ────────────────────────────────────────────────────────────

def build_report(df: pd.DataFrame, output_path: str):
    if df.empty:
        print("ERROR: No MoA data to report.")
        return

    stats = compute_statistics(df)

    print("Generating narrative with Gemini...")
    narrative = generate_executive_summary(stats, df)

    doc = Document()

    # ── Page setup ────────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    # ── Default font ──────────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.space_before = Pt(0)

    # ── Title ─────────────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(REPORT_TITLE)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(
        f"Generated {datetime.now().strftime('%B %d, %Y')}  •  "
        f"{stats['total_drugs']} Drug(s) Assessed"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    sub.paragraph_format.space_after = Pt(10)

    # Divider
    p_line = doc.add_paragraph()
    p_line.paragraph_format.space_after = Pt(6)
    pBdr = p_line._element.get_or_add_pPr().makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single", qn("w:sz"): "6",
        qn("w:space"): "1", qn("w:color"): "1F3864",
    })
    pBdr.append(bottom)
    p_line._element.get_or_add_pPr().append(pBdr)

    # ── Executive Summary ─────────────────────────────────────────────────────
    _styled_heading(doc, "Executive Summary")
    p = doc.add_paragraph(narrative.get("executive_summary", ""))
    p.paragraph_format.space_after = Pt(6)

    # ── Portfolio Overview ────────────────────────────────────────────────────
    _styled_heading(doc, "Portfolio Overview")

    overview_items = [
        ("Total Drugs Assessed", str(stats["total_drugs"])),
        ("Drugs Covered", ", ".join(stats["drugs"])),
        ("Average MoA Score", f"{stats['avg_score']} / 5" if stats["avg_score"] else "N/A"),
        ("Guardrail Summary", f"{stats['guardrail_pass']} PASS  |  {stats['guardrail_fail']} FAIL"),
    ]

    dist_parts = []
    for s in [5, 4, 3, 2, 1]:
        count = stats["score_distribution"].get(s, 0)
        if count > 0:
            dist_parts.append(f"{SCORE_LABEL[s]}: {count}")
    if dist_parts:
        overview_items.append(("Score Distribution", "; ".join(dist_parts)))

    cls_parts = []
    for cls_name, cnt in stats["classification_distribution"].items():
        cls_parts.append(f"{cls_name}: {cnt}")
    if cls_parts:
        overview_items.append(("Classification Distribution", "; ".join(cls_parts)))

    ov_table = doc.add_table(rows=len(overview_items), cols=2)
    ov_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    ov_table.style = "Table Grid"
    for i, (label, value) in enumerate(overview_items):
        cell_l = ov_table.rows[i].cells[0]
        cell_l.text = ""
        p = cell_l.paragraphs[0]
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(9)
        set_cell_shading(cell_l, "E8EDF3")

        cell_r = ov_table.rows[i].cells[1]
        cell_r.text = ""
        p = cell_r.paragraphs[0]
        run = p.add_run(value)
        run.font.size = Pt(9)

    for row_obj in ov_table.rows:
        row_obj.cells[0].width = Inches(2.5)
        row_obj.cells[1].width = Inches(4.2)

    doc.add_paragraph("")

    # ── Key Findings ──────────────────────────────────────────────────────────
    _styled_heading(doc, "Key Findings")
    for finding in narrative.get("key_findings", []):
        p = doc.add_paragraph(finding, style="List Bullet")
        p.paragraph_format.space_after = Pt(2)

    # ── Drug Score Summary Table ──────────────────────────────────────────────
    _styled_heading(doc, "MoA Innovation Score Summary")

    cols = ["drug_name", "indication", "moa_classification", "score", "guardrail",
            "confidence_tier", "analysis_date"]
    display_cols = ["Drug", "Indication", "Classification", "Score", "Guardrail",
                    "Confidence", "Date"]

    table = doc.add_table(rows=1, cols=len(cols))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, label in enumerate(display_cols):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "1F3864")

    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, col in enumerate(cols):
            val = row.get(col, "")
            cell = row_cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val) if pd.notna(val) else "N/A")
            run.font.size = Pt(8)

            if col == "score":
                try:
                    score_int = int(float(val))
                    if score_int in SCORE_COLOR_MAP:
                        run.font.color.rgb = SCORE_COLOR_MAP[score_int]
                        run.bold = True
                except (ValueError, TypeError):
                    pass
            elif col == "moa_classification":
                cls_str = str(val)
                for cls_key, cls_color in CLASSIFICATION_COLORS.items():
                    if cls_key.lower() in cls_str.lower():
                        run.font.color.rgb = cls_color
                        run.bold = True
                        break
            elif col == "guardrail":
                if str(val) == "PASS":
                    run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
                    run.bold = True
                elif str(val) == "FAIL":
                    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
                    run.bold = True

    widths = [Inches(1.1), Inches(1.2), Inches(1.1), Inches(0.5),
              Inches(0.7), Inches(0.8), Inches(0.8)]
    for row_obj in table.rows:
        for i, cell in enumerate(row_obj.cells):
            cell.width = widths[i]

    doc.add_paragraph("")

    # ── Classification Analysis ───────────────────────────────────────────────
    cls_text = narrative.get("classification_analysis", "")
    if cls_text:
        _styled_heading(doc, "Classification Analysis")
        doc.add_paragraph(cls_text)

    # ── Risk & Strength Highlights ────────────────────────────────────────────
    _styled_heading(doc, "Risk & Strength Analysis")

    p = doc.add_paragraph()
    run = p.add_run("Weak / At-Risk Drugs")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    doc.add_paragraph(narrative.get("risk_highlights", "N/A"))

    p = doc.add_paragraph()
    run = p.add_run("Strongest Innovators")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
    doc.add_paragraph(narrative.get("strength_highlights", "N/A"))

    # ── Per-Drug Breakdown ────────────────────────────────────────────────────
    per_drug = narrative.get("per_drug_narratives", {})
    if per_drug:
        _styled_heading(doc, "Drug-Level Analysis")

        for drug_name, drug_narrative in per_drug.items():
            h3 = doc.add_heading(drug_name, level=3)
            for run in h3.runs:
                run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
                run.font.size = Pt(11)

            drug_stat = stats.get("per_drug_stats", {}).get(drug_name, {})
            if drug_stat:
                stat_p = doc.add_paragraph()
                stat_run = stat_p.add_run(
                    f"Score: {drug_stat.get('score', 'N/A')}/5  |  "
                    f"Classification: {drug_stat.get('classification', 'N/A')}  |  "
                    f"Guardrail: {drug_stat.get('guardrail', 'N/A')}  |  "
                    f"Confidence: {drug_stat.get('confidence_tier', 'N/A')}"
                )
                stat_run.font.size = Pt(9)
                stat_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                stat_run.italic = True

            doc.add_paragraph(drug_narrative)

    # ── Scoring Framework ─────────────────────────────────────────────────────
    _styled_heading(doc, "MoA Innovation Scoring Framework")

    framework = [
        ("5", "Exceptional", "True First-in-Class: novel mechanism, strong rationale, class-creating potential"),
        ("4", "Strong", "Validated class with clearly superior mechanism differentiation (Best-in-Class)"),
        ("3", "Moderate", "Validated class with limited innovation (Me-too / Fast Follower)"),
        ("2", "Weak", "Older or strategically outdated mechanism"),
        ("1", "Poor", "Weak biological rationale or clinically invalidated mechanism"),
    ]

    sf_table = doc.add_table(rows=1, cols=3)
    sf_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sf_table.style = "Table Grid"
    for i, label in enumerate(["Score", "Label", "Description"]):
        cell = sf_table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "1F3864")

    for sc, lbl, desc in framework:
        row_cells = sf_table.add_row().cells
        row_cells[0].text = ""
        p = row_cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(sc)
        run.bold = True
        run.font.size = Pt(9)
        set_cell_shading(row_cells[0], "E8EDF3")

        row_cells[1].text = ""
        p = row_cells[1].paragraphs[0]
        run = p.add_run(lbl)
        run.bold = True
        run.font.size = Pt(9)

        row_cells[2].text = ""
        p = row_cells[2].paragraphs[0]
        p.add_run(desc).font.size = Pt(9)

    fw_widths = [Inches(0.5), Inches(1.2), Inches(4.6)]
    for row_obj in sf_table.rows:
        for i, cell in enumerate(row_obj.cells):
            cell.width = fw_widths[i]

    doc.add_paragraph("")

    # ── Score Legend ───────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    run = p.add_run("Score Legend: ")
    run.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    legend_text = "  |  ".join(f"{k} = {v}" for k, v in SCORE_LABEL.items())
    run = p.add_run(legend_text)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # ── Footer ────────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    run = p.add_run(
        "This report was auto-generated from MoA Innovation Scorer output "
        "using Gemini for narrative analysis."
    )
    run.font.size = Pt(7)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.save(output_path)
    print(f"\n✅ MoA report saved → {output_path}")


# ── GCS Upload ────────────────────────────────────────────────────────────────

def upload_to_gcs(local_path: str, drug_names: list) -> list:
    try:
        from google.cloud import storage
    except ImportError:
        raise ImportError("google-cloud-storage required. pip install google-cloud-storage")

    credentials = _get_credentials()
    client = storage.Client(project=BQ_PROJECT_ID, credentials=credentials)
    bucket = client.bucket(GCS_BUCKET)
    gcs_uris = []

    for drug_name in drug_names:
        safe_name = re.sub(r"[^a-zA-Z0-9_-]", "_", str(drug_name))
        blob_name = f"{GCS_BASE_PATH}/{safe_name}/{GCS_FILE_NAME}"
        gcs_uri = f"gs://{GCS_BUCKET}/{blob_name}"
        print(f"  Uploading to GCS: {gcs_uri}")
        try:
            blob = bucket.blob(blob_name)
            blob.upload_from_filename(
                local_path,
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            gcs_uris.append(gcs_uri)
        except Exception as e:
            print(f"  [ERROR] GCS upload failed for '{drug_name}': {e}")
            raise

    return gcs_uris


# ── Public entry point (for run_all.py) ───────────────────────────────────────

def generate_moa_report(molecule: str = None, output_path: str = None) -> str:
    """Generate MoA report. Returns the output file path."""
    if not API_KEY:
        print("[MoA Report] GEMINI_API_KEY not set — skipping report generation.")
        return ""

    if output_path is None:
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = f"moa_innovation_report_{ts}.docx"

    df = load_from_bigquery(molecule)
    if df.empty:
        print("[MoA Report] No data found — skipping.")
        return ""

    build_report(df, output_path)

    drug_names = df["drug_name"].dropna().unique().tolist()
    if drug_names:
        try:
            print(f"\nUploading MoA report to GCS for {len(drug_names)} drug(s)...")
            gcs_uris = upload_to_gcs(output_path, drug_names)
            for uri in gcs_uris:
                print(f"  ✅ {uri}")
        except Exception as e:
            print(f"  [WARN] GCS upload failed: {e}")

    return output_path


# ── CLI ───────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="Generate MoA Innovation Word report from BigQuery")
    parser.add_argument("--output", "-o", default=None, help="Output .docx path")
    parser.add_argument("--molecule", default=None, help="Filter to a specific molecule")
    args = parser.parse_args()

    generate_moa_report(molecule=args.molecule, output_path=args.output)


if __name__ == "__main__":
    main()
