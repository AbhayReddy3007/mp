import os
import re
import sys
import fitz
from docx import Document
from dotenv import load_dotenv
from google import genai
from google.genai import types

# ==============================
# LOAD ENV VARIABLES
# ==============================
load_dotenv()
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
if not GEMINI_API_KEY:
    raise ValueError("❌ GEMINI_API_KEY not found.")
client = genai.Client(api_key=GEMINI_API_KEY)

# ==============================
# CONFIG
# ==============================
INPUT_PDF  = "soc.pdf"
MODEL_NAME = "gemini-2.5-flash"

# ==============================
# PARSE CLI ARGUMENTS
# Usage: python code.py DrugA DrugB DrugC ...
# MOA is prompted interactively for each drug.
# ==============================
def parse_drugs() -> list[dict]:
    """
    Returns [{"name": "Cotadutide", "moa": "GLP1"}, ...]
    Drug names come from CLI args; MOA is asked interactively per drug.
    """
    if len(sys.argv) > 1:
        drug_names = sys.argv[1:]
    else:
        user_input = input(
            "No drugs supplied. Enter drug names separated by spaces: "
        ).strip()
        if not user_input:
            raise SystemExit("❌ No drug names provided. Exiting.")
        drug_names = user_input.split()

    print(f"\n💊 Drugs to analyse: {', '.join(drug_names)}")
    print("📋 Enter the MOA for each drug:\n")

    drugs = []
    for name in drug_names:
        moa = input(f"  MOA for {name}: ").strip()
        if not moa:
            raise SystemExit(f"❌ No MOA provided for {name}. Exiting.")
        drugs.append({"name": name, "moa": moa})

    print()
    return drugs

# ==============================
# SYSTEM INSTRUCTION  (SoC extraction)
# ==============================
SYSTEM_INSTRUCTION = """You are a Clinical Pharmacist and Diabetes/Obesity Treatment Pathway Analyst.
Analyze the provided Standard of Care (SoC) document and create a country-level Lines of Therapy (LOT) benchmark.

PURPOSE:
Neutral, evidence-based extraction for downstream LOT comparison.

SCOPE:
- Type 2 Diabetes Mellitus (T2DM)
- Obesity / Weight Management

OUTPUT FORMAT STRICT:
Country: [Country]
Country-Level SoC LOT Benchmark

1L:
Recommended therapies/interventions:
Patient segment or trigger:
Treatment goal or rationale:

2L:
Recommended therapies/interventions:
Patient segment or trigger:
Treatment goal or rationale:

3L:
Recommended therapies/interventions:
Patient segment or trigger:
Treatment goal or rationale:

Salvage:
Recommended therapies/interventions:
Patient segment or trigger:
Treatment goal or rationale:

Therapy Classes Explicitly Mentioned:
- List all major classes

FORMATTING RULES:
- Output plain business English.
- Do NOT use markdown.
- Do NOT use LaTeX.
- Use Unicode symbols directly: >= <= kg/m2.
- Do NOT use *, **, #, or $.
- Produce Word-document-ready text."""

# ==============================
# BUILD PER-DRUG OVERLAY PROMPT
# One prompt per drug = independent analysis each time
# ==============================
def build_drug_prompt(name: str, moa: str) -> str:
    return f"""You are a Senior Market Access Analyst.

Task:
Determine the most appropriate Line of Treatment (LOT) classification for the molecule below.
Base your decision ONLY on the supplied MOA and the SoC benchmark provided above.
Do NOT compare this molecule to any other drug. Analyse it independently on its own merits.

Molecule: {name}
MOA: {moa}

LOT ASSIGNMENT RULES:
- Identify the pharmacologic class, therapeutic modality, or treatment approach from the MOA.
- Match the MOA to the therapy classes, mechanisms, modalities, or treatment approaches explicitly described in the SoC.
- Determine the earliest treatment line in which the matching therapy class, mechanism, modality, or treatment approach appears.
- If the MOA aligns with multiple treatment lines, assign the earliest applicable line.
- Do NOT use historical treatment sequencing, current prescribing patterns, external guidelines, prior knowledge, or assumptions outside the supplied SoC.
- Evaluate the MOA against all SoC pathways and patient segments described in the benchmark.
- When a therapy class, mechanism, modality, or treatment approach appears in multiple SoC pathways, prioritize the earliest treatment line in which it is explicitly recommended.
- Only assign Second-Line when the MOA aligns primarily with therapies described as add-on, substitute, escalation, replacement, or post-failure options relative to First-Line treatment.
- Only assign Third-Line when the MOA aligns primarily with therapies described as later-line escalation, refractory-disease management, restricted-use therapies, or options used after failure of earlier treatment lines.
- Only assign Salvage when the MOA aligns primarily with rescue therapies, transplantation, last-resort interventions, or therapies explicitly described in the Salvage section of the SoC.
- When uncertainty exists, assign the earliest treatment line supported by the MOA-to-SoC mapping.
- In the Rationale, cite specific SoC lines and therapy classes from the benchmark to justify the assignment. Be specific to this molecule only.

CLASSIFICATION OPTIONS:
- First-line standard of care
- Strong first-line alternative / dominant second-line
- Second-line option
- Third-line or restricted niche use
- Salvage / last-resort use

OUTPUT FORMAT:
Molecule: {name}
MOA: {moa}
Final LoT Category:
Rationale:

FORMATTING RULES:
- Output plain business English.
- Do NOT use markdown.
- Do NOT use LaTeX.
- Use Unicode symbols directly.
- Produce Word-document-ready text."""

# ==============================
# PDF EXTRACTION
# ==============================
def extract_pdf_text(pdf_path: str) -> str:
    doc = fitz.open(pdf_path)
    text = "".join(page.get_text("text") for page in doc)
    doc.close()
    return text

# ==============================
# CLEAN GEMINI OUTPUT
# ==============================
def clean_text(text: str) -> str:
    replacements = {
        "$\\ge$": ">=", "$\\le$": "<=",
        "\\ge": ">=",   "\\le": "<=",
        "**": "", "*": "", "#": ""
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return re.sub(r"\$+", "", text).strip()

# ==============================
# GEMINI CALLS
# ==============================
def run_soc_extraction(pdf_text: str) -> str:
    response = client.models.generate_content(
        model=MODEL_NAME,
        contents=f"Below is the full SoC document:\n\n{pdf_text}",
        config=types.GenerateContentConfig(
            system_instruction=SYSTEM_INSTRUCTION,
            temperature=0.1
        )
    )
    return response.text or "No response received."

def run_drug_analysis(soc_text: str, name: str, moa: str) -> str:
    """
    Fully independent API call for a single drug.
    Sending only this drug's name + MOA ensures the model
    cannot copy-paste output from another drug in the same prompt.
    """
    prompt = build_drug_prompt(name, moa)
    response = client.models.generate_content(
        model=MODEL_NAME,
        contents=(
            "Below is the extracted SoC benchmark. "
            "Use it as the sole reference for LOT assignment.\n\n"
            f"{soc_text}\n\n{prompt}"
        ),
        config=types.GenerateContentConfig(temperature=0.1)
    )
    return response.text or "No response received."

# ==============================
# WRITE FORMATTED DOCX
# ==============================
def add_formatted_section(doc, text: str, drug_names: list[str]):
    drug_headings = {f"{d}:".lower() for d in drug_names}

    for line in text.split("\n"):
        line = clean_text(line)
        if not line:
            continue
        line_lower = line.lower()

        if line.startswith("Country:"):
            doc.add_heading(line, level=1)

        elif line in ["1L:", "2L:", "3L:", "Salvage:"]:
            doc.add_heading(line.replace(":", ""), level=2)

        elif (
            line_lower.startswith("recommended therapies")
            or line_lower.startswith("patient segment")
            or line_lower.startswith("treatment goal")
            or line_lower.startswith("molecule:")
            or line_lower.startswith("moa:")
            or line_lower.startswith("final lot category:")
            or line_lower.startswith("rationale:")
        ):
            p = doc.add_paragraph()
            p.add_run(line).bold = True

        elif line.startswith("-"):
            doc.add_paragraph(line[1:].strip(), style="List Bullet")

        elif line_lower in drug_headings:
            doc.add_heading(line.replace(":", ""), level=2)

        else:
            doc.add_paragraph(line)

# ==============================
# MAIN
# ==============================
def main():
    # 1. Collect drug names + MOAs from CLI / prompt
    drugs = parse_drugs()
    drug_names = [d["name"] for d in drugs]

    # 2. Dynamic output filename  e.g. soc_lot_cotadutide_mazdutide.docx
    slug = "_".join(d["name"].lower() for d in drugs)
    output_doc = f"soc_lot_{slug}.docx"

    # 3. Extract PDF text once
    print("📄 Reading PDF...")
    pdf_text = extract_pdf_text(INPUT_PDF)

    # 4. Extract SoC once (shared across all drug analyses)
    print("🧠 Extracting SoC...")
    soc_output = run_soc_extraction(pdf_text)

    # 5. One independent API call per drug
    #    Each call only knows about its own drug — no cross-contamination
    drug_results = {}
    for drug in drugs:
        print(f"⚡ Analysing {drug['name']} (MOA: {drug['moa']})...")
        drug_results[drug["name"]] = run_drug_analysis(
            soc_output, drug["name"], drug["moa"]
        )

    # 6. Build DOCX — SoC section first, then one page per drug
    print("📝 Generating DOCX...")
    doc = Document()

    doc.add_heading("Country-Level SoC LOT Benchmark", level=1)
    add_formatted_section(doc, soc_output, drug_names)

    for drug in drugs:
        doc.add_page_break()
        doc.add_heading(
            f"{drug['name']} — LOT Classification  |  MOA: {drug['moa']}",
            level=1
        )
        add_formatted_section(doc, drug_results[drug["name"]], drug_names)

    doc.save(output_doc)
    print(f"✅ Saved: {output_doc}")

if __name__ == "__main__":
    main()
